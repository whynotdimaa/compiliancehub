from io import BytesIO

import fitz
import pytest
from docx import Document as DocxDocument

from app.ingestion.parsing import ParsingError, parse_document


def test_markdown_heading_hierarchy():
    content = b"""# Policy
Intro line.
## Retention
Data is kept 5 years.
## Access
Only admins.
### Exceptions
Break-glass accounts.
"""
    sections = parse_document("policy.md", content)
    paths = [s.heading_path for s in sections]
    assert ("Policy",) in paths
    assert ("Policy", "Retention") in paths
    assert ("Policy", "Access", "Exceptions") in paths
    retention = next(s for s in sections if s.heading_path == ("Policy", "Retention"))
    assert "5 years" in retention.text


def test_markdown_sibling_heading_replaces_level():
    sections = parse_document("d.md", b"## A\ntext a\n## B\ntext b\n")
    assert [s.heading_path for s in sections] == [("A",), ("B",)]


def test_plain_text_single_section():
    sections = parse_document("notes.txt", b"Just some text.\nSecond line.")
    assert len(sections) == 1
    assert sections[0].heading_path == ()
    assert "Second line." in sections[0].text


def test_docx_headings_and_tables_in_order():
    docx = DocxDocument()
    docx.add_heading("Data Policy", level=1)
    docx.add_paragraph("Applies to all systems.")
    docx.add_heading("Retention", level=2)
    table = docx.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Type"
    table.cell(0, 1).text = "Period"
    table.cell(1, 0).text = "Logs"
    table.cell(1, 1).text = "1 year"
    buffer = BytesIO()
    docx.save(buffer)

    sections = parse_document("policy.docx", buffer.getvalue())
    top = next(s for s in sections if s.heading_path == ("Data Policy",))
    assert "Applies to all systems." in top.text
    retention = next(s for s in sections if s.heading_path == ("Data Policy", "Retention"))
    assert "Type | Period" in retention.text
    assert "Logs | 1 year" in retention.text


def test_pdf_font_size_heading_detection():
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Data Retention Policy", fontsize=20)
    page.insert_text((72, 110), "Records are kept for five years.", fontsize=11)
    page.insert_text((72, 130), "Backups are encrypted at rest.", fontsize=11)
    data = pdf.tobytes()

    sections = parse_document("policy.pdf", data)
    assert len(sections) == 1
    assert sections[0].heading_path == ("Data Retention Policy",)
    assert "five years" in sections[0].text
    assert sections[0].page == 1


def test_unsupported_extension_raises():
    with pytest.raises(ParsingError, match="Unsupported"):
        parse_document("malware.exe", b"MZ")


def test_broken_pdf_raises_parsing_error():
    with pytest.raises(ParsingError):
        parse_document("broken.pdf", b"not a pdf at all")
